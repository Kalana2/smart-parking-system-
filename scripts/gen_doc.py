from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ---- Title ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart Parking Management System\nFolder Structure & Implementation Plan')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(30, 60, 120)
doc.add_paragraph()

# ---- Helper functions ----
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 60, 120)

def add_code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.3)
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

# ---- 1. Overview ----
add_heading('1. Overview')
doc.add_paragraph(
    'This document describes the recommended folder structure for the Smart Parking Management System — '
    'an event-driven microservices application using Computer Vision (YOLOv8), RabbitMQ, Docker, '
    'and a React dashboard frontend.'
)
doc.add_paragraph(
    'The project uses a monorepo layout where each of the 10 microservices has its own directory, '
    'Dockerfile, dependencies, and tests — enabling independent development while keeping everything in one repository.'
)

# ---- 2. Root Structure ----
add_heading('2. Root Directory Structure')
add_code_block(
    'smart-parking-system/\n'
    '├── services/                    # All microservices\n'
    '├── frontend/                    # React.js dashboard\n'
    '├── shared/                      # Shared event contracts\n'
    '├── infra/                       # Docker & deployment configs\n'
    '├── docs/                        # Report, diagrams, API docs\n'
    '├── models/                      # ML model weights (.pt)\n'
    '├── scripts/                     # Dev/CI helper scripts\n'
    '├── docker-compose.yml           # Orchestrates all containers\n'
    '├── docker-compose.dev.yml       # Dev overrides\n'
    '├── .env.example                 # Environment variable template\n'
    '├── Makefile                     # Convenience commands\n'
    '└── README.md'
)

# ---- 3. Services Overview ----
add_heading('3. Microservices Overview')
add_table(
    ['Service', 'Language / Framework', 'Key Libraries'],
    [
        ['Detection + Tracking', 'Python 3.10+', 'YOLOv8, OpenCV, DeepSORT'],
        ['Slot Service', 'Python (FastAPI)', 'Shapely (IoU)'],
        ['ANPR Service', 'Python', 'YOLOv8, Tesseract OCR'],
        ['Dynamic Slot Service', 'Python (FastAPI)', 'OpenCV, NumPy'],
        ['Crash Detection', 'Python', 'OpenCV, Motion Analysis'],
        ['Verification Service', 'Node.js (Express)', 'Sequelize ORM'],
        ['Parking Aggregator', 'Node.js (NestJS)', 'Redis Client'],
        ['History Service', 'Node.js (Express)', 'Sequelize ORM'],
        ['API Gateway', 'Node.js (NestJS)', 'JWT, Express'],
        ['Web Dashboard', 'React.js', 'Axios, WebSocket'],
    ]
)

# ---- 4. Python Service ----
add_heading('4. Python Microservice Structure')
doc.add_paragraph('Each Python service (detection, slot, ANPR, dynamic-slot, crash) follows this layout:')
add_code_block(
    'services/detection-service/\n'
    '├── src/\n'
    '│   ├── __init__.py\n'
    '│   ├── main.py              # Entrypoint\n'
    '│   ├── config.py            # Env config (pydantic-settings)\n'
    '│   ├── detector.py          # YOLOv8 inference\n'
    '│   ├── tracker.py           # DeepSORT tracking\n'
    '│   ├── publisher.py         # RabbitMQ publisher\n'
    '│   ├── consumer.py          # RabbitMQ consumer\n'
    '│   └── utils/\n'
    '│       ├── bbox.py          # Bounding box helpers\n'
    '│       └── preprocessing.py # Frame preprocessing\n'
    '├── tests/\n'
    '├── Dockerfile\n'
    '├── requirements.txt\n'
    '├── .env.example\n'
    '└── README.md'
)
add_table(
    ['Service', 'Key Modules in src/'],
    [
        ['detection-service', 'detector.py, tracker.py, publisher.py'],
        ['slot-service', 'slot_manager.py, iou.py, consumer.py'],
        ['anpr-service', 'plate_detector.py, ocr.py, preprocessor.py'],
        ['dynamic-slot-service', 'gap_analyzer.py, classifier.py'],
        ['crash-detection-service', 'motion_analyzer.py, velocity_tracker.py'],
    ]
)

# ---- 5. Node.js Service ----
add_heading('5. Node.js Microservice Structure')
doc.add_paragraph('Each Node.js service (verification, aggregator, history, gateway) follows this layout:')
add_code_block(
    'services/verification-service/\n'
    '├── src/\n'
    '│   ├── index.js             # Entrypoint\n'
    '│   ├── app.js               # Express app setup\n'
    '│   ├── config/index.js      # Environment config\n'
    '│   ├── routes/vehicles.js   # REST endpoints\n'
    '│   ├── controllers/         # Request handlers\n'
    '│   ├── services/            # Business logic\n'
    '│   ├── models/Vehicle.js    # Sequelize model\n'
    '│   ├── messaging/\n'
    '│   │   ├── consumer.js      # RabbitMQ consumer\n'
    '│   │   └── publisher.js     # RabbitMQ publisher\n'
    '│   ├── middleware/\n'
    '│   └── utils/logger.js\n'
    '├── tests/\n'
    '├── Dockerfile\n'
    '├── package.json\n'
    '├── .env.example\n'
    '└── README.md'
)

# ---- 6. Frontend ----
add_heading('6. Frontend (React Dashboard)')
add_code_block(
    'frontend/\n'
    '├── public/index.html\n'
    '├── src/\n'
    '│   ├── App.js\n'
    '│   ├── api/parkingApi.js        # Axios client\n'
    '│   ├── components/\n'
    '│   │   ├── common/              # Buttons, Cards, Loaders\n'
    '│   │   ├── dashboard/           # Admin widgets\n'
    '│   │   ├── parking/             # Slot grid, map\n'
    '│   │   └── vehicles/            # Vehicle list, forms\n'
    '│   ├── pages/\n'
    '│   │   ├── AdminDashboard.jsx\n'
    '│   │   ├── UserDashboard.jsx\n'
    '│   │   ├── VehicleManagement.jsx\n'
    '│   │   └── History.jsx\n'
    '│   ├── hooks/useWebSocket.js\n'
    '│   ├── context/AuthContext.js\n'
    '│   └── styles/\n'
    '├── Dockerfile\n'
    '└── package.json'
)

# ---- 7. Shared ----
add_heading('7. Shared Contracts')
add_code_block(
    'shared/\n'
    '├── events/\n'
    '│   ├── vehicle.detected.v1.json\n'
    '│   ├── slot.updated.v1.json\n'
    '│   ├── plate.detected.v1.json\n'
    '│   ├── vehicle.verified.v1.json\n'
    '│   ├── dynamic.slot.detected.v1.json\n'
    '│   └── crash.detected.v1.json\n'
    '└── constants/event-names.js'
)

# ---- 8. Infra ----
add_heading('8. Infrastructure')
add_code_block(
    'infra/\n'
    '├── docker/\n'
    '│   ├── rabbitmq/rabbitmq.conf\n'
    '│   ├── postgres/init.sql\n'
    '│   └── redis/redis.conf\n'
    '├── nginx/nginx.conf\n'
    '└── scripts/\n'
    '    ├── seed-db.sh\n'
    '    └── wait-for-it.sh'
)

# ---- 9. Network Config ----
add_heading('9. Container Networking')
add_table(
    ['Service', 'Hostname', 'Port'],
    [
        ['RabbitMQ', 'rabbitmq', '5672 / 15672'],
        ['PostgreSQL', 'postgres', '5432'],
        ['Redis', 'redis', '6379'],
        ['API Gateway', 'gateway', '5000'],
        ['Frontend', 'frontend', '3000'],
    ]
)

# ---- 10. Root Files ----
add_heading('10. Root Configuration Files')
add_table(
    ['File', 'Purpose'],
    [
        ['docker-compose.yml', 'Production orchestration of all containers'],
        ['docker-compose.dev.yml', 'Dev overrides (hot reload, debug ports)'],
        ['.env.example', 'DB creds, RabbitMQ URL, Redis URL, model paths'],
        ['Makefile', 'Shortcuts: make up, make down, make logs, make test'],
        ['.gitignore', 'Ignore .env, node_modules, __pycache__, *.pt'],
    ]
)

# ---- Save ----
out = '/home/snake/Projects/Smart Parking System/docs/Smart_Parking_Folder_Structure.docx'
doc.save(out)
print(f'✅ Document saved to: {out}')
